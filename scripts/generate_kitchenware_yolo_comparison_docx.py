"""
One-off / utility: build Word report with side-by-side comparison.
Run: uv run python scripts/generate_kitchenware_yolo_comparison_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "artifacts" / "reports" / "kitchenware-florence2-vs-yolo-coco-comparison.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Kitchenware (Florence-2) vs Phase 1 COCO (YOLO11)\nSide-by-Side Evaluation Report"
    )
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph("Robotic vision project — summary of runs and metrics.").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph()

    doc.add_heading("1. What we did (Kitchenware / Kaggle)", level=1)
    for line in [
        "Data: Kaggle Kitchenware Classification (train.csv + images). Six classes: cup, fork, glass, knife, plate, spoon.",
        "Adaptation: Built a benchmark manifest with one full-image bounding box per image so Florence-2 could be fine-tuned on the existing <OD> detection training path.",
        "Splits: Stratified train / val / test_real_heldout (seed 42). Full manifest: 5,559 labeled images.",
        "Model: microsoft/Florence-2-base, 1 epoch, AdamW 1e-5, grad accum 8, FP16, CUDA.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2. Side-by-side: overall metrics (held-out test)", level=1)
    doc.add_paragraph(
        "Compare the headline numbers from each run’s evaluation JSON. "
        "Prefer mAP50 when comparing to Florence-2 (mAP50-95 is duplicated in the Florence report for this pipeline)."
    )

    t1 = doc.add_table(rows=5, cols=3)
    t1.style = "Table Grid"
    hdr = t1.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Florence-2 — Kitchenware (Kaggle)"
    hdr[2].text = "YOLO11 — Phase 1 COCO benchmark"
    for c in hdr:
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True

    rows_data = [
        ("mAP50", "0.368", "0.453"),
        ("mAP50-95", "0.368 *", "0.303"),
        ("Precision", "0.409", "0.491"),
        ("Recall", "0.612", "0.470"),
    ]
    for i, (m, f, y) in enumerate(rows_data, start=1):
        t1.rows[i].cells[0].text = m
        t1.rows[i].cells[1].text = f
        t1.rows[i].cells[2].text = y

    doc.add_paragraph()
    doc.add_paragraph(
        "* Florence-2 evaluation report stores the same value for mAP50 and mAP50-95 in this codebase (see train_florence2 evaluation)."
    ).italic = True

    doc.add_heading("3. Side-by-side: per-class metrics", level=1)
    doc.add_paragraph(
        "Kitchenware: mAP50 per class (Florence-2). "
        "COCO Phase 1: mAP50-95 per class from YOLO11 report (two classes: mug, book). "
        "These label sets are different — not directly comparable class-for-class."
    )

    t2 = doc.add_table(rows=7, cols=4)
    t2.style = "Table Grid"
    h2 = t2.rows[0].cells
    h2[0].text = "Kitchenware class"
    h2[1].text = "mAP50"
    h2[2].text = "COCO Phase 1 class"
    h2[3].text = "mAP50-95"
    for c in h2:
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True

    kw = [
        ("cup", "0.825"),
        ("plate", "0.524"),
        ("glass", "0.483"),
        ("knife", "0.156"),
        ("fork", "0.129"),
        ("spoon", "0.093"),
    ]
    coco = [("mug", "0.478"), ("book", "0.128")]
    for i in range(6):
        row = t2.rows[i + 1].cells
        row[0].text = kw[i][0]
        row[1].text = kw[i][1]
        if i < len(coco):
            row[2].text = coco[i][0]
            row[3].text = coco[i][1]
        else:
            row[2].text = "—"
            row[3].text = "—"

    doc.add_heading("4. Which performed better (summary)", level=1)
    for line in [
        "mAP50 and precision: YOLO11 on the Phase 1 COCO benchmark is higher (0.453 vs 0.368; precision 0.491 vs 0.409).",
        "Recall: Florence-2 on kitchenware is higher (0.612 vs 0.470).",
        "Why not a single winner: different number of classes (6 vs 2), different supervision (full-image pseudo-boxes vs real COCO boxes), and YOLO trained many epochs on a dedicated detector task vs one epoch Florence-2 fine-tuning.",
        "YOLO is stronger on strict detection AP for its two-class COCO subset; Florence-2 shows strong signal on some kitchenware classes (e.g. cup) under the adapted setup.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("5. File references in this repo", level=1)
    t3 = doc.add_table(rows=5, cols=2)
    t3.style = "Table Grid"
    refs = [
        ("Kitchenware Florence-2 report", "artifacts/reports/kitchenware-kaggle-florence2-full.json"),
        ("Kitchenware manifest", "artifacts/manifests/kitchenware-kaggle-full.json"),
        ("YOLO11 best eval", "artifacts/reports/yolo11-best-eval.json"),
        ("Phase 1 YOLO brief", "artifacts/reports/phase1-summary-brief.md"),
        ("Markdown version of this report", "artifacts/reports/kitchenware-florence2-vs-yolo-coco-report.md"),
    ]
    for i, (a, b) in enumerate(refs):
        t3.rows[i].cells[0].text = a
        t3.rows[i].cells[1].text = b

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
